import os
import io
import re
import zipfile
from datetime import datetime
import pandas as pd
import streamlit as st

# ---- Soft imports (si faltan, no rompen) ----
try:
    import docx as _docx  # paquete: python-docx
except Exception:
    _docx = None

try:
    import pdfplumber as _pdfplumber
except Exception:
    _pdfplumber = None

try:
    import pytesseract as _pytesseract
    from pdf2image import convert_from_bytes as _convert_from_bytes
except Exception:
    _pytesseract = None
    _convert_from_bytes = None

# ---- Utilidades ----
def normalize_text(t: str) -> str:
    t = re.sub(r"[ \t]+", " ", t)
    t = re.sub(r"\n{2,}", "\n", t).strip()
    return t

def find_after(label: str, text: str, maxlen: int = 400, flags=re.IGNORECASE) -> str:
    pattern = rf"{re.escape(label)}\s*[:\-]?\s*(.+)"
    m = re.search(pattern, text, flags)
    if m:
        val = m.group(1).strip().split("\n")[0].strip()
        return val[:maxlen]
    return ""

def to_int_num(s: str | None):
    if not s:
        return None
    s = re.sub(r"[^\d]", "", s)
    return int(s) if s.isdigit() else None

# ---- Extractores ----
def docx_to_text(file_like) -> str:
    if _docx is not None:
        try:
            document = _docx.Document(file_like)
            parts = [p.text for p in document.paragraphs]
            for t in document.tables:
                for row in t.rows:
                    parts.append(" | ".join([cell.text.strip() for cell in row.cells]))
            return "\n".join(parts)
        except Exception as e:
            st.warning(f"No se pudo leer DOCX con python-docx: {e}")
    try:
        raw = file_like.read() if hasattr(file_like, "read") else file_like
        with zipfile.ZipFile(io.BytesIO(raw)) as z:
            xml = z.read("word/document.xml").decode("utf-8", errors="ignore")
        xml = re.sub(r"</w:p>", "\n", xml)
        text = re.sub(r"<[^>]+>", "", xml)
        text = re.sub(r"\s+\n", "\n", text)
        return text
    except Exception as e:
        st.error(f"No se pudo leer el DOCX (fallback ZIP/XML): {e}")
        return ""

def pdf_to_text(raw_bytes: bytes) -> str:
    if _pdfplumber is not None:
        try:
            with _pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
                pages = [page.extract_text() or "" for page in pdf.pages]
                txt = "\n".join(pages)
                if txt.strip():
                    return txt
        except Exception as e:
            st.warning(f"pdfplumber falló: {e}")
    if _pytesseract is not None and _convert_from_bytes is not None:
        try:
            images = _convert_from_bytes(raw_bytes, dpi=300)
            out = []
            for img in images:
                out.append(_pytesseract.image_to_string(img, lang="spa+eng"))
            return "\n".join(out)
        except Exception as e:
            st.warning(f"OCR falló: {e}")
            return ""
    return ""

def extract_text_any(file_bytes: bytes, ext: str):
    logs = []
    if ext == "docx":
        text = docx_to_text(io.BytesIO(file_bytes))
        logs.append("DOCX leído (python-docx o fallback ZIP/XML).")
        if not text.strip():
            logs.append("No se pudo extraer texto del DOCX.")
        return text, logs

    if ext == "pdf":
        if _pdfplumber is not None:
            logs.append("Intentando PDF digital con pdfplumber…")
        else:
            logs.append("pdfplumber no disponible; intentaré OCR…")

        if _pytesseract is not None and _convert_from_bytes is not None:
            logs.append("OCR disponible (pytesseract + pdf2image).")
        else:
            logs.append("OCR NO disponible. Revisa requirements.txt y packages.txt.")

        text = pdf_to_text(file_bytes)
        if text.strip():
            logs.append("Texto extraído (digital u OCR).")
            return text, logs
        else:
            logs.append("No se pudo extraer texto (ni pdfplumber ni OCR).")
            return "", logs

    return "", ["Extensión no soportada."]

# ---- Parseo de campos ----
def extract_fields(text: str) -> dict:
    t = normalize_text(text)
    m_date = re.search(r"(\d{1,2}\s+de\s+\w+\s+de\s+\d{4})", t, re.IGNORECASE)
    fecha_doc = m_date.group(1) if m_date else ""
    nombre = find_after("NOMBRE", t)
    unidad = find_after("REQUIRENTE (UNIDAD)", t)
    objetivo = find_after("OBJETIVO", t)
    monto_est = to_int_num(find_after("MONTO ESTIMADO", t))
    return {
        "Fecha Documento": fecha_doc,
        "Solicitante (Nombre)": nombre,
        "Unidad Requirente": unidad,
        "Objetivo": objetivo,
        "Monto Estimado": monto_est,
    }

# ---- App ----
REG_PATH = os.path.join("data", "registro_solicitudes.csv")
os.makedirs(os.path.dirname(REG_PATH), exist_ok=True)

st.set_page_config(page_title="Registro Solicitudes de Compra – SLEP Petorca", layout="wide")
st.title("📄 Registro de Solicitudes de Compra – V°B° Financiero SLEP Petorca")

with st.expander("🔧 Diagnóstico del entorno", expanded=False):
    import shutil, subprocess
    st.write(f"python-docx: **{_docx is not None}**")
    st.write(f"pdfplumber: **{_pdfplumber is not None}**")
    st.write(f"OCR (pytesseract+pdf2image): **{_pytesseract is not None and _convert_from_bytes is not None}**")
    st.write("tesseract:", shutil.which("tesseract") or "NO ENCONTRADO")
    st.write("pdftoppm:", shutil.which("pdftoppm") or "NO ENCONTRADO")

tab1, tab2 = st.tabs(["➕ Ingresar solicitud", "📊 Registro y Resumen"])

with tab1:
    up = st.file_uploader("Sube la solicitud (DOCX o PDF)", type=["docx", "pdf"])
    fecha_recepcion = st.date_input("Fecha de recepción", datetime.now().date())
    firmado = st.checkbox("Marcar como Firmado ahora")
    fecha_firma = st.date_input("Fecha firma V°B°", datetime.now().date()) if firmado else None

    if up:
        ext = up.name.lower().split(".")[-1]
        raw = up.read()
        with st.status("Extrayendo texto…", expanded=False):
            text, logs = extract_text_any(raw, ext)
            for line in logs:
                st.write("• " + line)

        if not text.strip():
            st.error("No se pudo extraer texto. Si es PDF escaneado, revisa packages.txt.")
        else:
            if st.checkbox("Ver texto extraído"):
                st.text_area("Contenido", value=text, height=250)
            data = extract_fields(text)
            st.json(data)
            if st.button("Guardar en registro"):
                df = pd.read_csv(REG_PATH) if os.path.exists(REG_PATH) else pd.DataFrame()
                row = {
                    **data,
                    "Fecha de Recepción": fecha_recepcion.strftime("%d/%m/%Y"),
                    "Fecha Firma V°B°": fecha_firma.strftime("%d/%m/%Y") if firmado else "",
                    "Estado": "Firmado" if firmado else "Pendiente",
                    "Archivo Origen": up.name,
                }
                df = pd.concat([df, pd.DataFrame([row])], ignore_index=True)
                df.to_csv(REG_PATH, index=False)
                st.success("Registro guardado correctamente.")

with tab2:
    if os.path.exists(REG_PATH):
        df = pd.read_csv(REG_PATH)
        st.dataframe(df, use_container_width=True, hide_index=True)
        st.download_button("⬇️ Descargar CSV", data=df.to_csv(index=False).encode("utf-8"),
                           file_name="registro_solicitudes.csv", mime="text/csv")
    else:
        st.info("Aún no hay registro guardado.")
